"""Формирование проектов документов: акт проверки и предписание.

Юридически значимых действий система не совершает: формируется проект,
который инспектор проверяет, при необходимости правит и подписывает.
В документ попадают только подтверждённые инспектором гипотезы.
"""
from __future__ import annotations

import io
from datetime import date

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt

from analysis.models import Finding

SEVERITY_RU = {"critical": "критическое", "major": "существенное", "minor": "незначительное"}


def _base_document(title: str, region: dict) -> Document:
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(11)
    for section in doc.sections:
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(1.5)

    header = doc.add_paragraph(region.get("supervisory_body",
                                          "Комитет государственного строительного надзора "
                                          "города Москвы"))
    header.alignment = WD_ALIGN_PARAGRAPH.CENTER
    header.runs[0].bold = True

    heading = doc.add_paragraph(title)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    heading.runs[0].bold = True
    return doc


def _requisites(doc: Document, card: dict, inspector: dict, run_id: str) -> None:
    table = doc.add_table(rows=0, cols=2)
    table.style = "Table Grid"
    rows = [
        ("Объект капитального строительства", card.get("name", "")),
        ("Адрес", card.get("address", "")),
        ("Разрешение на строительство", f"№ {card.get('permit_number', '')} от "
                                        f"{card.get('permit_date', '')}"),
        ("Застройщик", card.get("developer", "")),
        ("Лицо, осуществляющее строительство", card.get("contractor", "")),
        ("Должностное лицо", f"{inspector.get('position', '')} {inspector.get('full_name', '')}"),
        ("Дата составления", date.today().strftime("%d.%m.%Y")),
        ("Идентификатор анализа", run_id),
    ]
    for name, value in rows:
        cells = table.add_row().cells
        cells[0].text = name
        cells[1].text = str(value)


def _finding_block(doc: Document, index: int, finding: Finding) -> None:
    para = doc.add_paragraph()
    run = para.add_run(f"{index}. {finding.title}")
    run.bold = True

    element = finding.structural_element or {}
    location = ", ".join(p for p in [element.get("name", ""), element.get("mark", ""),
                                     f"оси {element['axes']}" if element.get("axes") else "",
                                     f"отм. {element['level']}" if element.get("level") else "",
                                     element.get("section", "")] if p and p != "—")
    doc.add_paragraph(f"Конструктивный элемент: {location or 'не определён'}")
    doc.add_paragraph(f"Классификация: {finding.code}, {SEVERITY_RU.get(finding.severity, '')} "
                      f"расхождение, уверенность модели {finding.confidence:.0%}")

    norms = "; ".join(f"{r['doc']} п. {r['clause']} (ред. {r['edition']}"
                      + (", действующая)" if r.get("actual") else ", проверить актуальность)")
                      for r in finding.norm_refs)
    doc.add_paragraph(f"Нормативное основание: {norms}")
    if finding.legal_refs:
        legal = "; ".join(f"{r['act']} ст. {r['article']} ч. {r['part']}"
                          for r in finding.legal_refs)
        doc.add_paragraph(f"Правовое основание: {legal}")

    doc.add_paragraph("Доказательства:")
    for ev in finding.evidence:
        bbox = ", ".join(f"{v:.0f}" for v in ev.bbox)
        doc.add_paragraph(f"— {ev.doc_title}, лист {ev.page}, область [{bbox}]: "
                          f"{ev.extracted} ({ev.role})", style="List Bullet")
    if finding.field_check:
        doc.add_paragraph(f"Проверить на объекте: {finding.field_check}")
    if finding.documents_to_request:
        doc.add_paragraph("Истребовать документы: " + "; ".join(finding.documents_to_request))
    doc.add_paragraph()


def _signature(doc: Document, inspector: dict) -> None:
    doc.add_paragraph()
    para = doc.add_paragraph(f"{inspector.get('position', 'Инспектор')}"
                             f"\t\t____________________\t\t{inspector.get('full_name', '')}")
    para.alignment = WD_ALIGN_PARAGRAPH.LEFT


def build_inspection_act(card: dict, findings: list[Finding], inspector: dict,
                         run_id: str, region: dict | None = None) -> bytes:
    """Проект акта проверки по подтверждённым гипотезам."""
    region = region or {}
    doc = _base_document("ПРОЕКТ АКТА ПРОВЕРКИ", region)
    _requisites(doc, card, inspector, run_id)

    doc.add_paragraph()
    doc.add_paragraph("Основание проверки: программа проверок объекта капитального "
                      "строительства, статья 54 Градостроительного кодекса Российской Федерации.")
    doc.add_paragraph("В ходе анализа представленной документации выявлены следующие "
                      "расхождения, подтверждённые должностным лицом:")
    doc.add_paragraph()
    for index, finding in enumerate(findings, start=1):
        _finding_block(doc, index, finding)

    doc.add_paragraph("Настоящий документ сформирован автоматически как проект и подлежит "
                      "проверке должностным лицом. Выводы системы носят характер гипотез "
                      "и не являются заключением до их подтверждения.")
    _signature(doc, inspector)
    return _to_bytes(doc)


def build_prescription(card: dict, findings: list[Finding], inspector: dict, run_id: str,
                       deadline: date, region: dict | None = None) -> bytes:
    """Проект предписания об устранении выявленных нарушений."""
    region = region or {}
    doc = _base_document("ПРОЕКТ ПРЕДПИСАНИЯ ОБ УСТРАНЕНИИ НАРУШЕНИЙ", region)
    _requisites(doc, card, inspector, run_id)

    doc.add_paragraph()
    doc.add_paragraph(f"{card.get('contractor', 'Лицу, осуществляющему строительство')} "
                      "предписывается устранить следующие нарушения "
                      f"в срок до {deadline.strftime('%d.%m.%Y')}:")
    doc.add_paragraph()
    for index, finding in enumerate(findings, start=1):
        _finding_block(doc, index, finding)

    doc.add_paragraph("О выполнении настоящего предписания уведомить орган государственного "
                      "строительного надзора с приложением подтверждающих документов.")
    _signature(doc, inspector)
    return _to_bytes(doc)


def _to_bytes(doc: Document) -> bytes:
    buffer = io.BytesIO()
    doc.save(buffer)
    return buffer.getvalue()
